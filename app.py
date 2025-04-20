from flask import Flask, request, jsonify, render_template, send_file
import os
import openai
import google.generativeai as genai
import chromadb
from langchain_openai import OpenAIEmbeddings
import psycopg2
import json
from dotenv import load_dotenv
from docx import Document
import random

# ==== Flask 設定 ====
app = Flask(__name__)

# ==== 載入 .env ====
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
DATABASE_URL = os.getenv("DATABASE_URL")

openai_client = openai.OpenAI(api_key=OPENAI_API_KEY)
genai.configure(api_key=GEMINI_API_KEY)

# ==== 向量資料庫 (備用學系特色檢索) ====
embeddings = OpenAIEmbeddings(api_key=OPENAI_API_KEY)
db = chromadb.PersistentClient(path="./chroma_db")
collection = db.get_or_create_collection(name="department_info")

# ==== 連接 PostgreSQL 資料庫 ====
def get_db_connection():
    return psycopg2.connect(DATABASE_URL)

# ==== 代碼對應表 ====
CODE_MAPPING = {
    "A": "修課紀錄", "B": "書面報告", "C": "實作作品", "D": "自然科學探究",
    "E": "社會領域探究", "F": "高中自主學習計畫", "G": "社團活動經驗", "H": "幹部經驗",
    "I": "服務學習經驗", "J": "競賽表現", "K": "非修課成果作品", "L": "檢定證照",
    "M": "特殊優良表現", "N": "多元表現綜整心得", "O": "高中學習歷程反思",
    "P": "學習動機", "Q": "未來學習計畫"
}

# ==== 分項表與Prompt ====
SECTION_MAPPING = {
    "ABC": {  
        "codes": ["A", "B", "C", "D", "E"],
        "prompt": "請撰寫一篇課程學習成果，強調學習歷程與探索精神，並與學系關聯性相結合。"
    },
    "N": {  
        "codes": ["F", "G", "H", "I", "J", "K", "L", "M"],
        "prompt": "請了解學系專業與選材理念後，撰寫強調自主學習的多元表現綜整心得，包含引言、3個段落(其中每個段落需有一個標題，總結該段落的內容)，以及總結，並利用使用者輸入的具體事例舉例每件事的能力成長與省思。內容與學系密切相關，且必須確保字數在800字內。。"
    },
    "O": {  
        "codes": ["A", "B", "C", "D", "E", "F", "G", "H", "I", "J", "K", "L", "M"],
        "prompt": "請撰寫高中學習歷程反思，強調學習經驗、挑戰與成長，並舉出失敗與反省以展現學習態度。必須確保字數在800字內。"
    },
    "P": {  
        "codes": ["A", "B", "C", "D", "E", "F", "G", "H", "I", "J", "K", "L", "M"],
        "prompt": "請撰寫學習動機，強調與學系高度相關的興趣與企圖心，並適當佐證。",
        "needs_department_features": True
    },
    "Q": {  
        "codes": ["A", "B", "C", "D", "E", "F", "G", "H", "I", "J", "K", "L", "M"],
        "prompt": "請撰寫未來計畫，分成三個階段：(1)入學前想學習的課程與技能 (2)就讀期間計畫修習的課程 (3)畢業後的職業發展。",
        "needs_department_features": True
    }
}

# ==== 風格Prompt ====
STYLE_PROMPTS = {
    "formal": "語氣自然、像高中生在對學長姐或教授分享經驗。句子可以稍口語、真誠、有情緒，但仍保持基本邏輯與清晰度。避免誇張詞彙與過度文學。",
    "casual": "請以自然、口語化的語氣撰寫，像是在向教授講述自己的故事。內容可稍微白話，但仍應保持邏輯與專業度。",
    "concise": "請以簡潔明快的語言撰寫，避免冗詞贅句，直接切入重點，強調邏輯清晰與重點聚焦。",
    "detailed": "請以完整、深入的方式撰寫，內容須具體，包含明確例子與說明，讓讀者能全面理解申請者的能力與動機。",
}

# ==== QA ====
@app.route("/ask", methods=["POST"])
def ask_question():
    data = request.get_json()
    university = data.get("university")
    department = data.get("department")
    question = data.get("question")

    if not university or not department or not question:
        return jsonify({"answer": "請提供完整的學校、學系與問題內容！"}), 400

    # 取得該學系特色與需繳交資料作為回答基礎
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT section_code, section_name, content 
        FROM application_guidelines 
        WHERE university = %s AND department = %s
    """, (university, department))
    sections = cursor.fetchall()
    cursor.close()
    conn.close()

    department_features = get_department_features(university, department)

    base_context = "\n".join([
        f"{code} - {name}：{content}" for code, name, content in sections
    ]) or "查無備審項目內容。"

    prompt = f"""
你是一名大學備審客服人員，請針對使用者提出的問題，根據下列內容進行回答：
- **請務必將每一條條列點獨立寫在新的一行（請用 \n 分隔），不要將多個條列點寫在同一行。**


【學校】{university}
【學系】{department}
【系所特色】{department_features}
【備審資料】{base_context}

使用者提問：
{question}

請以簡潔清楚的語氣回覆，字數不超過150字，必要時可引用學系特色或備審資料，但不得亂編。
    """

    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
        )
        answer = response.choices[0].message.content.strip()
        return jsonify({"answer": answer})
    except Exception as e:
        return jsonify({"answer": f"⚠️ 回覆錯誤：{str(e)}"}), 500



# ==== 提取學系特色 (優先從資料庫，其次從向量資料庫) ====
def get_department_features(university, department):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT features FROM department_features 
        WHERE university = %s AND department = %s
    """, (university, department))
    result = cursor.fetchone()
    
    cursor.close()
    conn.close()

    if result:
        return result[0]  # 如果資料庫有學系特色，則回傳
    
    # 如果資料庫無結果，則查詢向量資料庫
    vector_search = collection.query(query_texts=[f"{university} {department} 學系特色"], n_results=1)
    if "documents" in vector_search and vector_search["documents"]:
        return vector_search["documents"][0]
    
    return "未提供學系特色"

# ==== 查詢需繳交資料與學系特色（首頁） ====
@app.route("/check", methods=["GET", "POST"])
def check_requirements():
    results = None
    department_features = None
    university = department = ""

    if request.method == "POST":
        university = request.form.get("university")
        department = request.form.get("department")

        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT section_code, section_name FROM application_guidelines WHERE university=%s AND department=%s", (university, department))
        results = cursor.fetchall()
        department_features = get_department_features(university, department)
        cursor.close()
        conn.close()

    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT DISTINCT university, department FROM application_guidelines")
    all_options = cursor.fetchall()
    cursor.close()
    conn.close()

    return render_template(
        "check_requirements.html",
        all_options=all_options,
        results=results,
        department_features=department_features,
        university=university,
        department=department
    )

# ==== 獲取所有學校與學系(check) ====
@app.route("/get_departments", methods=["GET"])
def get_departments_by_school():
    university = request.args.get("university")
    if not university:
        return jsonify([])

    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT DISTINCT department 
        FROM application_guidelines 
        WHERE university = %s
    """, (university,))
    departments = [row[0] for row in cursor.fetchall()]
    cursor.close()
    conn.close()
    
    return jsonify(departments)

# ==== 獲取所有學校與學系(generator) ====
@app.route("/get_universities", methods=["GET"])
def get_universities():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT DISTINCT university FROM application_guidelines")
    universities = [row[0] for row in cursor.fetchall()]
    cursor.close()
    conn.close()
    return jsonify(universities)


# ==== 學系備審指引查詢 ====
@app.route("/get_guidelines", methods=["POST"])
def get_guidelines():
    data = request.json
    university = data.get("university")
    department = data.get("department")

    if not university or not department:
        return jsonify({"error": "請選擇學校與學系"}), 400

    conn = get_db_connection()
    cursor = conn.cursor()

    # 查詢 application_guidelines
    cursor.execute("""
        SELECT section_code, section_name, content 
        FROM application_guidelines 
        WHERE university=%s AND department=%s
    """, (university, department))
    results = cursor.fetchall()

    # 查詢 department_features
    department_features = get_department_features(university, department)

    cursor.close()
    conn.close()

    if not results:
        return jsonify({"error": "該學校學系尚無備審資料"}), 400

    sections = {row[0]: {"name": row[1], "content": row[2]} for row in results}

    return jsonify({
        "university": university,
        "department": department,
        "sections": sections,
        "department_features": department_features
    })

# ==== 首頁 ====
@app.route("/")
def index():
    return render_template("index.html")

# ==== 備審段落生成器 ====
@app.route("/generate", methods=["GET"])
def generate_page():
    return render_template("generator.html")

# ==== 查詢備審指引頁（如果有）====
@app.route("/show_check_page", methods=["GET", "POST"])
def show_check_page():
    # 處理查詢資料邏輯
    return render_template("check_requirements.html")

# ==== 生成段落 ====
@app.route("/generate_paragraph", methods=["POST"])
def generate_paragraph():
    data = request.get_json()
    
    section_code = data.get("section_code")
    user_inputs = data.get("user_inputs", {})
    university = data.get("university", "未知大學")
    department = data.get("department", "未知學系")
    style = data.get("style", "formal")  
    word_count = random.randint(600, 800) 
    adjust_percentage = random.randint(30, 50) 
    
    if section_code not in SECTION_MAPPING:
        return jsonify({"error": "無效的段落代碼"}), 400
    
     # 獲取學系特色
    department_features = get_department_features(university, department)
        
    # 獲取撰寫指引 (section_prompt)
    section_prompt = SECTION_MAPPING[section_code]["prompt"]
    
    input_text = "\n".join(f"{CODE_MAPPING.get(code, code)}: {text}" for code, text in user_inputs.items() if text)

    # 使用LLM初次生成
    first_prompt = f"""
    你是一名高中生，正在申請{university}-{department}，請根據以下內容撰寫 {section_prompt}。

    請務必遵守以下規則，嚴格按照使用者提供的內容進行撰寫，不得編造資訊：
    1.只能使用以下學系特色，不得新增額外內容：{department_features}
    2.只能使用使用者提供的內容，不得自行發揮：{input_text}
    3.字數範圍：{word_count} 字以內
    4.風格要求：{STYLE_PROMPTS.get(style, STYLE_PROMPTS['formal'])}
    5.不得杜撰經歷、不得添加未提供的競賽、活動、學術研究等內容。
    6.內容應清晰、邏輯合理、段落流暢，並忠實呈現使用者輸入的重點。

    請根據這些要求，產生一段符合申請需求的內容。
    """
    try:
        gemini_model = genai.GenerativeModel("gemini-2.0-flash")
        gemini_response = gemini_model.generate_content(
            first_prompt, 
            generation_config={"temperature": 0.6}
        )
        first_output = gemini_response.text.strip()
    except Exception as e:
        return jsonify({"error": f"Gemini 生成內容時發生錯誤: {str(e)}"}), 500

    # 進行 N 次優化
    N = 1
    for _ in range(N):
        refine_prompt = f"""
        這是先前生成的內容：
        {first_output}

        請根據以下要求進一步優化：
        1. 請調整語句，使表達方式稍有不同，但仍然保持相同核心內容
        2. 請確保字數範圍在 {word_count} 字以內
        3. 請保持 {STYLE_PROMPTS.get(style, STYLE_PROMPTS['formal'])} 的語氣
        4. 請讓語句更流暢，避免過於生硬
        5.只能使用以下學系特色，不得新增額外內容：{department_features}
        6.只能使用使用者提供的內容，不得自行發揮：{input_text}
        7.變更幅度：約 {adjust_percentage}%

        """
        # 動態調整 temperature
        temperature_value = 0.3 + (adjust_percentage / 100) * 0.4
        if adjust_percentage > 30:
            refine_prompt += "\n請嘗試替換一些詞語，使表達方式更加生動。"
        if adjust_percentage > 50:
            refine_prompt += "\n請嘗試變換句子結構，使內容更加流暢自然。"
        if adjust_percentage > 70:
            refine_prompt += "\n請嘗試用不同的方式表達相同意思，使表達方式多樣化。"
        
        try:
            gemini_model = genai.GenerativeModel("gemini-2.0-flash")
            gemini_response = gemini_model.generate_content(
                refine_prompt, 
                generation_config={"temperature": temperature_value}
            )
            final_output = gemini_response.text.strip()
        except Exception as e:
            return jsonify({"error": f"Gemini 生成內容時發生錯誤: {str(e)}"}), 500
        
    # 語氣優化
    improved_prompt = f"""
    你是一位擅長潤飾大學申請備審資料的語言專家，請協助我修改以下備審資料，使其自然流暢，語氣真誠，並保留原有內容重點。

    目前的內容：
    {final_output}

    請確保：
    1.不可新增資訊，只能在原始內容基礎上進行潤飾與語句調整
    2.風格請採用：{STYLE_PROMPTS.get(style, STYLE_PROMPTS['formal'])}
    3.請避免使用以下 AI 常見用語或句型：如「總體而言」、「本文將探討」、「綜上所述」、「在當今社會中」、「產生深遠影響」、「我堅信」等
    4.請避免過於工整、生硬、過度學術化的句式結構，讓整體語氣更貼近一位真誠且有思考力的高中學生口吻
    5.每個新段落請用 \n\n 分隔，輸出為純文字，不使用 Markdown 語法
    6.請將文字濃縮至 {word_count} 字以內，保留關鍵資訊與主要邏輯，刪除重複或不必要的詞語，使內容更精練但不失自然語感。

    優化後的內容：
    """
    try:    
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": improved_prompt}],
            temperature=0.5,
            #max_tokens=int(word_count * 1.3),
            stop=["優化後的內容："]
        )
        improved_output = response.choices[0].message.content.strip()
    except Exception as e:
        return jsonify({"error": f"GPT 生成內容時發生錯誤: {str(e)}"}), 500
    
    return jsonify({
        "style": style,
        "adjust_percentage": adjust_percentage,
        "generated_text": improved_output
    }), 200, {'Content-Type': 'application/json; charset=utf-8'}

# ==== 生成 Word 文件 ====
@app.route("/generate_docx", methods=["POST"])
def generate_docx():
    data = request.get_json()
    university = data.get("university", "未知大學")
    department = data.get("department", "未知學系")
    section_code = data.get("section_code", "未知段落")
    section_name = CODE_MAPPING.get(section_code, section_code)
    generated_text = data.get("generated_text", "")

    if not generated_text:
        return jsonify({"error": "沒有生成內容"}), 400

    doc = Document()
    doc.add_heading(f"{university}-{department}", level=1)
    doc.add_heading(f"{section_code}{section_name}", level=2)
    doc.add_paragraph(generated_text)
    download_name = f"{university}{department}_{section_code}.docx"

    file_path = f"/tmp/{download_name}"
    doc.save(file_path)

    return send_file(file_path, as_attachment=True, download_name=f"{university}{department}_{section_code}.docx")

if __name__ == "__main__":
    app.run(debug=True)
